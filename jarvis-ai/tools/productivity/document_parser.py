"""
JARVIS Document Parser - Multi-format document extraction
=========================================================

Parses PDF, DOCX, TXT, MD, and JSON files for RAG ingestion.
"""

import os
from pathlib import Path
from typing import Tuple, Dict, Any, Optional
from datetime import datetime
import json


class DocumentParser:
    """
    Multi-format document parser.
    
    Supports:
    - PDF (requires PyMuPDF)
    - DOCX (requires python-docx)
    - TXT, MD, JSON (built-in)
    """
    
    def __init__(self):
        """Initialize parser with available backends."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which parsing libraries are available."""
        self.pdf_available = False
        self.docx_available = False
        
        try:
            import fitz  # PyMuPDF
            self.pdf_available = True
        except ImportError:
            pass
        
        try:
            import docx
            self.docx_available = True
        except ImportError:
            pass
    
    def parse(self, filepath: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a document and extract text content.
        
        Args:
            filepath: Path to the document
            
        Returns:
            Tuple of (content, metadata)
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        suffix = filepath.suffix.lower()
        
        # Get basic metadata
        metadata = {
            "filename": filepath.name,
            "filepath": str(filepath.absolute()),
            "extension": suffix,
            "size_bytes": filepath.stat().st_size,
            "modified_at": datetime.fromtimestamp(
                filepath.stat().st_mtime
            ).isoformat()
        }
        
        # Parse based on file type
        if suffix == '.pdf':
            content = self._parse_pdf(filepath)
        elif suffix == '.docx':
            content = self._parse_docx(filepath)
        elif suffix in ['.txt', '.md', '.markdown']:
            content = self._parse_text(filepath)
        elif suffix == '.json':
            content = self._parse_json(filepath)
        elif suffix in ['.html', '.htm']:
            content = self._parse_html(filepath)
        elif suffix in ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs']:
            content = self._parse_code(filepath)
        else:
            # Try as text
            try:
                content = self._parse_text(filepath)
            except Exception:
                content = ""
        
        return content, metadata
    
    def _parse_pdf(self, filepath: Path) -> str:
        """Parse PDF file."""
        if not self.pdf_available:
            raise ImportError(
                "PyMuPDF not installed. Run: pip install PyMuPDF"
            )
        
        import fitz
        
        text_parts = []
        
        with fitz.open(filepath) as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, filepath: Path) -> str:
        """Parse DOCX file."""
        if not self.docx_available:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )
        
        import docx
        
        doc = docx.Document(filepath)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_text(self, filepath: Path) -> str:
        """Parse plain text file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return filepath.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # Last resort: read with errors ignored
        return filepath.read_text(encoding='utf-8', errors='ignore')
    
    def _parse_json(self, filepath: Path) -> str:
        """Parse JSON file into readable text."""
        content = filepath.read_text(encoding='utf-8')
        
        try:
            data = json.loads(content)
            # Convert to readable format
            return self._json_to_text(data)
        except json.JSONDecodeError:
            return content
    
    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert JSON data to readable text."""
        lines = []
        prefix = "  " * indent
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_text(value, indent + 1))
                else:
                    lines.append(f"{prefix}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{i}]:")
                    lines.append(self._json_to_text(item, indent + 1))
                else:
                    lines.append(f"{prefix}- {item}")
        else:
            lines.append(f"{prefix}{data}")
        
        return "\n".join(lines)
    
    def _parse_html(self, filepath: Path) -> str:
        """Parse HTML file, extracting text content."""
        content = filepath.read_text(encoding='utf-8', errors='ignore')
        
        try:
            from html.parser import HTMLParser
            
            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_parts = []
                    self.skip_tags = {'script', 'style', 'nav', 'footer', 'header'}
                    self.current_tag = None
                
                def handle_starttag(self, tag, attrs):
                    self.current_tag = tag
                
                def handle_endtag(self, tag):
                    if tag in ['p', 'div', 'br', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self.text_parts.append('\n')
                    self.current_tag = None
                
                def handle_data(self, data):
                    if self.current_tag not in self.skip_tags:
                        text = data.strip()
                        if text:
                            self.text_parts.append(text)
            
            extractor = TextExtractor()
            extractor.feed(content)
            return ' '.join(extractor.text_parts)
        except Exception:
            # Fallback: strip HTML tags with regex
            import re
            text = re.sub(r'<[^>]+>', ' ', content)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
    
    def _parse_code(self, filepath: Path) -> str:
        """Parse code file, preserving docstrings and comments."""
        content = filepath.read_text(encoding='utf-8', errors='ignore')
        
        # Add filename as context
        header = f"# File: {filepath.name}\n"
        return header + content


# Convenience functions
def parse_document(filepath: str) -> Tuple[str, Dict[str, Any]]:
    """Parse a document and return content with metadata."""
    parser = DocumentParser()
    return parser.parse(filepath)


def extract_text(filepath: str) -> str:
    """Extract text content from a document."""
    content, _ = parse_document(filepath)
    return content
